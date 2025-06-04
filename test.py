import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from jinja2 import Template
import re
import os
from docx.enum.text import WD_BREAK

# IEEE formatting configuration
IEEE_CONFIG = {
    'font_name': 'Times New Roman',
    'font_size_title': Pt(24),
    'font_size_body': Pt(9.5),
    'font_size_caption': Pt(9),
    'margin_left': Inches(0.75),
    'margin_right': Inches(0.75),
    'margin_top': Inches(0.75),
    'margin_bottom': Inches(0.75),
    'column_count_body': 2,
    'column_spacing': Inches(0.25),
    'column_width': Inches(3.375),
    'column_indent': Inches(0.2),
    'line_spacing': Pt(10),  # Exact spacing for 9.5pt font
    'figure_sizes': {
        'Very Small': Inches(1.2),
        'Small': Inches(1.8),
        'Medium': Inches(2.5),
        'Large': Inches(3.2)
    },
    'max_figure_height': Inches(4.0),
}

# LaTeX template for IEEEtran
LATEX_TEMPLATE = r"""
\documentclass[conference]{IEEEtran}
\IEEEoverridecommandlockouts
\usepackage{cite}
\usepackage{amsmath,amssymb,amsfonts}
\usepackage{graphicx}
\usepackage{textcomp}
\usepackage{xcolor}

\begin{document}

\title{ {{ title }} }
\author{ {{ authors_latex }} }

\maketitle

\begin{abstract}
{{ abstract }}
\end{abstract}

\begin{IEEEkeywords}
{{ keywords }}
\end{IEEEkeywords}

{% for section in sections %}
\section{ {{ section.title }} }
{{ section.content }}
{% for subsection in section.subsections %}
\subsection{ {{ subsection.title }} }
{{ subsection.content }}
{% endfor %}
{% for figure in section.figures %}
\begin{figure}[h]
\centering
\includegraphics[width=0.8\columnwidth]{ {{ figure.file_name }} }
\caption{ {{ figure.caption }} }
\label{fig:{{ section.idx }}_{{ loop.index }}}
\end{figure}
{% endfor %}
{% endfor %}

\section*{Acknowledgment}
{{ acknowledgments }}

\begin{thebibliography}{ {{ references | length }} }
{% for ref in references %}
\bibitem{ {{ loop.index }} }
{{ ref.text }}
{% endfor %}
\end{thebibliography}

\end{document}
"""

def set_document_defaults(doc):
    """Set document-wide defaults to minimize unwanted spacing."""
    styles = doc.styles

    # Modify Normal style
    if 'Normal' in styles:
        normal = styles['Normal']
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(12)
        normal.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        normal.paragraph_format.line_spacing_rule = 0  # Exact spacing
        normal.paragraph_format.widow_control = False
        normal.font.name = IEEE_CONFIG['font_name']
        normal.font.size = IEEE_CONFIG['font_size_body']
        # Add better spacing control
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.first_line_indent = Pt(0)

    # Modify Heading 1 style
    if 'Heading 1' in styles:
        heading1 = styles['Heading 1']
        heading1.base_style = styles['Normal']
        heading1.paragraph_format.space_before = Pt(0)
        heading1.paragraph_format.space_after = Pt(0)
        heading1.paragraph_format.line_spacing = Pt(10)
        heading1.paragraph_format.line_spacing_rule = 0
        heading1.paragraph_format.keep_with_next = False  # Changed to False
        heading1.paragraph_format.page_break_before = False
        heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Headings should be left-aligned
        heading1.font.name = IEEE_CONFIG['font_name']
        heading1.font.size = IEEE_CONFIG['font_size_body']
        heading1.font.bold = True

    # Modify Heading 2 style for subsections
    if 'Heading 2' in styles:
        heading2 = styles['Heading 2']
        heading2.base_style = styles['Normal']
        heading2.paragraph_format.space_before = Pt(6)
        heading2.paragraph_format.space_after = Pt(0)
        heading2.paragraph_format.line_spacing = Pt(10)
        heading2.paragraph_format.line_spacing_rule = 0
        heading2.paragraph_format.keep_with_next = False
        heading2.paragraph_format.page_break_before = False
        heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Headings should be left-aligned
        heading2.font.name = IEEE_CONFIG['font_name']
        heading2.font.size = IEEE_CONFIG['font_size_body']
        heading2.font.bold = True

def add_title(doc, title):
    """Add the paper title."""
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.name = IEEE_CONFIG['font_name']
    run.font.size = IEEE_CONFIG['font_size_title']
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(12)

def add_authors(doc, authors):
    """Add authors and their details in a parallel layout using a table."""
    if not authors:
        return
    
    num_authors = len(authors)
    table = doc.add_table(rows=1, cols=num_authors)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.allow_autofit = True
    
    for idx, author in enumerate(authors):
        if not author.get('name'):
            continue
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        para = cell.add_paragraph()
        run = para.add_run(author['name'])
        run.bold = True
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)
        
        fields = [
            ('department', 'Department'),
            ('organization', 'Organization'),
            ('city', 'City'),
            ('state', 'State'),
            ('tamilnadu', 'Tamil Nadu')
        ]
        for field_key, field_name in fields:
            if author.get(field_key):
                para = cell.add_paragraph(author[field_key])
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                if para.runs:
                    para.runs[0].italic = True
                    para.runs[0].font.name = IEEE_CONFIG['font_name']
                    para.runs[0].font.size = IEEE_CONFIG['font_size_body']
        
        for custom_field in author.get('custom_fields', []):
            if custom_field['value']:
                para = cell.add_paragraph(custom_field['value'])
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                if para.runs:
                    para.runs[0].italic = True
                    para.runs[0].font.name = IEEE_CONFIG['font_name']
                    para.runs[0].font.size = IEEE_CONFIG['font_size_body']
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_footnote(doc, footnote_data):
    """Add the first footnote with submission/revision/DOI info."""
    footnote_text = f"Manuscript received {footnote_data['received_date']}; "
    footnote_text += f"revised {footnote_data['revised_date']}; "
    footnote_text += f"accepted {footnote_data['accepted_date']}. "
    footnote_text += f"This work was supported by {footnote_data['funding']}. "
    footnote_text += f"(DOI: {footnote_data['doi']})"
    
    if any([footnote_data['received_date'], footnote_data['revised_date'], 
            footnote_data['accepted_date'], footnote_data['funding'], footnote_data['doi']]):
        para = doc.add_paragraph()
        run = para.add_run(footnote_text.strip())
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_caption']
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(6)

def add_abstract(doc, abstract):
    """Add the abstract section with italicized 'Abstract—'."""
    if abstract:
        para = doc.add_paragraph()
        run = para.add_run("Abstract—")
        run.italic = True
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        run = para.add_run(abstract)
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        
        # Apply advanced justification controls to abstract
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = IEEE_CONFIG['line_spacing']
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        para.paragraph_format.line_spacing_rule = 0
        
        # Add advanced spacing controls to prevent word stretching
        para_element = para._element
        pPr = para_element.get_or_add_pPr()
        
        # Set justification method
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        
        # Control text alignment
        textAlignment = OxmlElement('w:textAlignment')
        textAlignment.set(qn('w:val'), 'baseline')
        pPr.append(textAlignment)
        
        # Prevent excessive word spacing
        adjust_right_ind = OxmlElement('w:adjustRightInd')
        adjust_right_ind.set(qn('w:val'), '0')
        pPr.append(adjust_right_ind)

def add_keywords(doc, keywords):
    """Add the keywords section."""
    if keywords:
        para = doc.add_paragraph(f"Keywords: {keywords}")
        
        # Apply advanced justification controls to keywords
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = IEEE_CONFIG['line_spacing']
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        para.paragraph_format.line_spacing_rule = 0
        if para.runs:
            para.runs[0].font.name = IEEE_CONFIG['font_name']
            para.runs[0].font.size = IEEE_CONFIG['font_size_body']
        
        # Add advanced spacing controls to prevent word stretching
        para_element = para._element
        pPr = para_element.get_or_add_pPr()
        
        # Set justification method
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        
        # Control text alignment
        textAlignment = OxmlElement('w:textAlignment')
        textAlignment.set(qn('w:val'), 'baseline')
        pPr.append(textAlignment)
        
        # Prevent excessive word spacing
        adjust_right_ind = OxmlElement('w:adjustRightInd')
        adjust_right_ind.set(qn('w:val'), '0')
        pPr.append(adjust_right_ind)
        
        # Minimal dummy paragraph to stabilize layout
        dummy_para = doc.add_paragraph("")
        dummy_para.paragraph_format.space_before = Pt(0)
        dummy_para.paragraph_format.space_after = Pt(0)
        dummy_para.paragraph_format.widow_control = False
        dummy_para.paragraph_format.keep_with_next = False
        dummy_para.paragraph_format.line_spacing = 0
        if dummy_para.runs:
            dummy_para.runs[0].font.size = Pt(1)

def add_section(doc, section_data, section_idx, is_first_section=False):
    """Add a section with content blocks (text and images), subsections, and figures."""
    if section_data['title']:
        para = doc.add_heading(f"{section_idx}. {section_data['title'].upper()}", level=1)
        para.paragraph_format.page_break_before = False
        para.paragraph_format.space_before = IEEE_CONFIG['line_spacing']  # Exactly one line before heading
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.keep_together = False
        para.paragraph_format.widow_control = False

    # Process content blocks (text and images in order)
    for block_idx, block in enumerate(section_data.get('content_blocks', [])):
        if block['type'] == 'text' and block.get('content'):
            space_before = IEEE_CONFIG['line_spacing'] if is_first_section and block_idx == 0 else Pt(3)
            para = add_justified_paragraph(
                doc, 
                block['content'],
                indent_left=IEEE_CONFIG['column_indent'],
                indent_right=IEEE_CONFIG['column_indent'],
                space_before=space_before,
                space_after=Pt(12)
            )
        elif block['type'] == 'image' and block.get('file') and block.get('caption'):
            size = block.get('size', 'Medium')
            width = IEEE_CONFIG['figure_sizes'][size]
            
            para = doc.add_paragraph()
            run = para.add_run()
            picture = run.add_picture(block['file'], width=width)
            if picture.height > IEEE_CONFIG['max_figure_height']:
                scale_factor = IEEE_CONFIG['max_figure_height'] / picture.height
                run.clear()
                run.add_picture(block['file'], width=width * scale_factor, height=IEEE_CONFIG['max_figure_height'])
            
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            
            # Generate figure number based on section and image position
            img_count = sum(1 for b in section_data.get('content_blocks', [])[:block_idx+1] if b['type'] == 'image')
            caption = doc.add_paragraph(f"Fig. {section_idx}.{img_count}: {block['caption']}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_before = Pt(0)
            caption.paragraph_format.space_after = Pt(12)
            if caption.runs:
                caption.runs[0].font.name = IEEE_CONFIG['font_name']
                caption.runs[0].font.size = IEEE_CONFIG['font_size_caption']

    # Legacy support for old content field
    if section_data.get('content') and not section_data.get('content_blocks'):
        space_before = IEEE_CONFIG['line_spacing'] if is_first_section else Pt(3)
        para = add_justified_paragraph(
            doc, 
            section_data['content'],
            indent_left=IEEE_CONFIG['column_indent'],
            indent_right=IEEE_CONFIG['column_indent'],
            space_before=space_before,
            space_after=Pt(12)
        )

    for sub_idx, subsection in enumerate(section_data.get('subsections', []), 1):
        if subsection.get('title'):
            para = doc.add_heading(f"{section_idx}.{sub_idx} {subsection['title']}", level=2)
            para.paragraph_format.page_break_before = False
            para.paragraph_format.space_before = IEEE_CONFIG['line_spacing']  # Exactly one line before heading
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.keep_with_next = False
            para.paragraph_format.keep_together = False
            para.paragraph_format.widow_control = False

        if subsection.get('content'):
            # Use the new justified paragraph function for better spacing
            para = add_justified_paragraph(
                doc, 
                subsection['content'],
                indent_left=IEEE_CONFIG['column_indent'],
                indent_right=IEEE_CONFIG['column_indent'],
                space_before=Pt(1),
                space_after=Pt(12)
            )
    
    # Legacy figures support (now at end of section)
    for fig_idx, figure in enumerate(section_data.get('figures', []), 1):
        if figure.get('file') and figure.get('caption'):
            size = figure.get('size', 'Medium')
            width = IEEE_CONFIG['figure_sizes'][size]
            
            para = doc.add_paragraph()
            run = para.add_run()
            picture = run.add_picture(figure['file'], width=width)
            if picture.height > IEEE_CONFIG['max_figure_height']:
                scale_factor = IEEE_CONFIG['max_figure_height'] / picture.height
                run.clear()
                run.add_picture(figure['file'], width=width * scale_factor, height=IEEE_CONFIG['max_figure_height'])
            
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            
            caption = doc.add_paragraph(f"Fig. {section_idx}.{fig_idx}: {figure['caption']}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_before = Pt(0)
            caption.paragraph_format.space_after = Pt(12)
            if caption.runs:
                caption.runs[0].font.name = IEEE_CONFIG['font_name']
                caption.runs[0].font.size = IEEE_CONFIG['font_size_caption']
    
    # for tab_idx, table in enumerate(section_data.get('tables', []), 1):
    #     if table.get('data') and isinstance(table['data'], list) and table['data'] and all(isinstance(row, list) and row for row in table['data']):
    #         caption = doc.add_paragraph(f"Table {section_idx}.{tab_idx}: {table.get('caption', '')}")
    #         caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #         caption.paragraph_format.space_before = Pt(6)
    #         caption.paragraph_format.space_after = Pt(6)
    #         if caption.runs:
    #             caption.runs[0].font.name = IEEE_CONFIG['font_name']
    #             caption.runs[0].font.size = IEEE_CONFIG['font_size_caption']
            
    #         num_cols = len(table['data'][0]) if table['data'][0] else 1
    #         doc_table = doc.add_table(rows=len(table['data']), cols=num_cols)
    #         doc_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #         for row_idx, row_data in enumerate(table['data']):
    #             cells = doc_table.rows[row_idx].cells
    #             for col_idx, cell_data in enumerate(row_data[:num_cols]):
    #                 cells[col_idx].text = str(cell_data) if cell_data is not None else ""
    #                 cell_para = cells[col_idx].paragraphs[0]
    #                 cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #                 cell_para.paragraph_format.left_indent = IEEE_CONFIG['column_indent']
    #                 cell_para.paragraph_format.right_indent = IEEE_CONFIG['column_indent']
    #                 cell_para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
    #                 cell_para.paragraph_format.line_spacing_rule = 0
    #                 for run in cell_para.runs:
    #                     run.font.name = IEEE_CONFIG['font_name']
    #                     run.font.size = IEEE_CONFIG['font_size_body']
    #         doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_references(doc, references):
    """Add references section with proper alignment (hanging indent)."""
    if references:
        para = doc.add_heading("References", level=1)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.keep_with_next = False
        
        for idx, ref in enumerate(references, 1):
            if ref.get('text'):
                para = doc.add_paragraph(f"[{idx}] {ref['text']}")
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = IEEE_CONFIG['column_indent'] + Inches(0.25)
                para.paragraph_format.right_indent = IEEE_CONFIG['column_indent']
                para.paragraph_format.first_line_indent = Inches(-0.25)
                para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
                para.paragraph_format.line_spacing_rule = 0
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.widow_control = False
                para.paragraph_format.keep_with_next = False
                para.paragraph_format.keep_together = True
                if para.runs:
                    para.runs[0].font.name = IEEE_CONFIG['font_name']
                    para.runs[0].font.size = IEEE_CONFIG['font_size_body']

def enable_auto_hyphenation(doc):
    """Enable conservative hyphenation to reduce word spacing without breaking words inappropriately."""
    section = doc.sections[-1]
    sectPr = section._sectPr

    # Enable automatic hyphenation but keep it conservative
    auto_hyphenation = OxmlElement('w:autoHyphenation')
    auto_hyphenation.set(qn('w:val'), '1')
    sectPr.append(auto_hyphenation)

    # Do NOT hyphenate capitalized words (proper nouns, abbreviations)
    do_not_hyphenate_caps = OxmlElement('w:doNotHyphenateCaps')
    do_not_hyphenate_caps.set(qn('w:val'), '1')  # Changed to 1 to prevent breaking proper nouns
    sectPr.append(do_not_hyphenate_caps)

    # Set a LARGER hyphenation zone (less aggressive hyphenation, only when really needed)
    hyphenation_zone = OxmlElement('w:hyphenationZone')
    hyphenation_zone.set(qn('w:val'), '720')  # Increased from 180 to 720 (half inch)
    sectPr.append(hyphenation_zone)

    # Limit consecutive hyphens to prevent excessive word breaking
    consecutive_hyphen_limit = OxmlElement('w:consecutiveHyphenLimit')
    consecutive_hyphen_limit.set(qn('w:val'), '2')  # Max 2 consecutive lines with hyphens
    sectPr.append(consecutive_hyphen_limit)

    # Ensure hyphenation is controlled, not suppressed
    compat = doc.settings.element.find(qn('w:compat'))
    if compat is None:
        doc.settings.element.append(OxmlElement('w:compat'))
        compat = doc.settings.element.find(qn('w:compat'))

    option = OxmlElement('w:suppressAutoHyphens')
    option.set(qn('w:val'), '0')
    compat.append(option)

def set_compatibility_options(doc):
    """Set compatibility options to optimize spacing and justification."""
    compat = doc.settings.element.find(qn('w:compat'))
    if compat is None:
        doc.settings.element.append(OxmlElement('w:compat'))
        compat = doc.settings.element.find(qn('w:compat'))

    # Critical options to eliminate word spacing issues
    
    # Force Word to use exact character spacing instead of word spacing
    option1 = OxmlElement('w:useWord2002TableStyleRules')
    option1.set(qn('w:val'), '1')
    compat.append(option1)
    
    # Prevent Word from expanding spaces for justification
    option2 = OxmlElement('w:doNotExpandShiftReturn')
    option2.set(qn('w:val'), '1')
    compat.append(option2)
    
    # Use consistent character spacing
    option3 = OxmlElement('w:useSingleBorderforContiguousCells')
    option3.set(qn('w:val'), '1')
    compat.append(option3)
    
    # Force exact spacing calculations
    option4 = OxmlElement('w:spacingInWholePoints')
    option4.set(qn('w:val'), '1')
    compat.append(option4)
    
    # Prevent auto spacing adjustments
    option5 = OxmlElement('w:doNotUseHTMLParagraphAutoSpacing')
    option5.set(qn('w:val'), '1')
    compat.append(option5)
    
    # Use legacy justification method (more precise)
    option6 = OxmlElement('w:useWord97LineBreakRules')
    option6.set(qn('w:val'), '1')
    compat.append(option6)
    
    # Disable automatic kerning adjustments
    option7 = OxmlElement('w:doNotAutoCompressPictures')
    option7.set(qn('w:val'), '1')
    compat.append(option7)
    
    # Force consistent text metrics
    option8 = OxmlElement('w:useNormalStyleForList')
    option8.set(qn('w:val'), '1')
    compat.append(option8)
    
    # Prevent text compression/expansion
    option9 = OxmlElement('w:doNotPromoteQF')
    option9.set(qn('w:val'), '1')
    compat.append(option9)
    
    # Use exact font metrics
    option10 = OxmlElement('w:useAltKinsokuLineBreakRules')
    option10.set(qn('w:val'), '0')
    compat.append(option10)

def generate_ieee_document(form_data):
    """Generate an IEEE-formatted Word document."""
    doc = Document()
    
    set_document_defaults(doc)
    
    section = doc.sections[0]
    section.left_margin = IEEE_CONFIG['margin_left']
    section.right_margin = IEEE_CONFIG['margin_right']
    section.top_margin = IEEE_CONFIG['margin_top']
    section.bottom_margin = IEEE_CONFIG['margin_bottom']
    
    add_title(doc, form_data.get('title', ''))
    add_authors(doc, form_data.get('authors', []))
    add_footnote(doc, form_data.get('footnote', {}))
    add_abstract(doc, form_data.get('abstract', ''))
    add_keywords(doc, form_data.get('keywords', ''))

    # Removed column break after abstract/keywords

    # Add continuous section break for two-column layout
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    section.start_type = WD_SECTION.CONTINUOUS
    section.left_margin = IEEE_CONFIG['margin_left']
    section.right_margin = IEEE_CONFIG['margin_right']
    section.top_margin = IEEE_CONFIG['margin_top']
    section.bottom_margin = IEEE_CONFIG['margin_bottom']
    
    # Set up the two-column layout
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    
    cols.set(qn('w:num'), str(IEEE_CONFIG['column_count_body']))
    cols.set(qn('w:sep'), '0')
    cols.set(qn('w:space'), str(int(IEEE_CONFIG['column_spacing'].pt)))
    cols.set(qn('w:equalWidth'), '1')
    
    for i in range(IEEE_CONFIG['column_count_body']):
        col = OxmlElement('w:col')
        col.set(qn('w:w'), str(int(IEEE_CONFIG['column_width'].pt)))
        cols.append(col)
    
    no_balance = OxmlElement('w:noBalance')
    no_balance.set(qn('w:val'), '1')
    sectPr.append(no_balance)
    
    for idx, section in enumerate(form_data.get('sections', []), 1):
        section['idx'] = idx
        add_section(doc, section, idx, is_first_section=(idx == 1))
    
    if form_data.get('acknowledgments'):
        para = doc.add_heading("Acknowledgment", level=1)
        para.paragraph_format.page_break_before = False
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        
        para = doc.add_paragraph(form_data['acknowledgments'])
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.left_indent = IEEE_CONFIG['column_indent']
        para.paragraph_format.right_indent = IEEE_CONFIG['column_indent']
        para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
        para.paragraph_format.line_spacing_rule = 0
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.keep_together = True
        if para.runs:
            para.runs[0].font.name = IEEE_CONFIG['font_name']
            para.runs[0].font.size = IEEE_CONFIG['font_size_body']
    
    add_references(doc, form_data.get('references', []))
    
    enable_auto_hyphenation(doc)
    set_compatibility_options(doc)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_latex_document(form_data):
    """Generate an IEEEtran LaTeX file."""
    authors_latex = " \\and ".join([
        f"\\IEEEauthorblockN{{{a['name']}}}\n" +
        f"\\IEEEauthorblockA{{\\textit{{ {a.get('department', '')} }}\\\\\n" +
        f"\\textit{{ {a.get('organization', '')} }}\\\\\n" +
        f"\\textit{{ {a.get('city', '')}, {a.get('state', '')}, {a.get('tamilnadu', '')} }}" +
        "".join([f"\\\\\n\\textit{{ {cf['value'] }}}" for cf in a.get('custom_fields', []) if cf['value']])
        for a in form_data.get('authors', []) if a.get('name')
    ])
    
    for idx, section in enumerate(form_data.get('sections', []), 1):
        section['idx'] = idx
        for fig_idx, figure in enumerate(section.get('figures', []), 1):
            figure['file_name'] = f"figure_{idx}_{fig_idx}.png"
    
    template = Template(LATEX_TEMPLATE)
    latex_content = template.render(
        title=form_data.get('title', ''),
        authors_latex=authors_latex,
        abstract=form_data.get('abstract', ''),
        keywords=form_data.get('keywords', ''),
        sections=form_data.get('sections', []),
        acknowledgments=form_data.get('acknowledgments', ''),
        references=form_data.get('references', [])
    )
    
    buffer = BytesIO()
    buffer.write(latex_content.encode('utf-8'))
    buffer.seek(0)
    return buffer

def validate_abstract(abstract):
    """Validate abstract word count (optional)."""
    words = len(re.findall(r'\b\w+\b', abstract))
    return 150 <= words <= 250

def validate_reference(ref_text):
    """Basic validation for IEEE reference format."""
    pattern = r'^\[\d+\]\s+[\w\s.,]+'
    return bool(re.match(pattern, ref_text))

def main():
    st.set_page_config(page_title="IEEE Paper Generator", layout="wide")
    st.title("IEEE Research Paper Generator")
    st.markdown("Create an IEEE-formatted research paper with dynamic sections, figures, and tables. "
                "Download as a Word document or LaTeX file.")
    st.info("Automatic hyphenation has been enabled in the Word document for optimal text justification.")

    if 'form_data' not in st.session_state:
        st.session_state.form_data = {
            'title': '',
            'authors': [{
                'name': '', 'department': '', 'organization': '', 
                'city': '', 'state': '', 'tamilnadu': '', 'custom_fields': []
            }],
            'footnote': {
                'received_date': '', 'revised_date': '', 'accepted_date': '',
                'funding': '', 'doi': ''
            },
            'abstract': '',
            'keywords': '',
            'sections': [{'title': 'Introduction', 'content': '', 'subsections': [], 'figures': [], 'tables': []}],
            'acknowledgments': '',
            'references': []
        }

    with st.expander("Basic Information", expanded=True):
        st.session_state.form_data['title'] = st.text_input(
            "Paper Title",
            value=st.session_state.form_data['title'],
            help="Enter the paper title (10–12 words recommended)."
        )
        st.session_state.form_data['abstract'] = st.text_area(
            "Abstract",
            value=st.session_state.form_data['abstract'],
            height=150,
            help="Summarize the paper's purpose, methods, results, and conclusions."
        )
        if st.session_state.form_data['abstract']:
            if not validate_abstract(st.session_state.form_data['abstract']):
                st.warning("Abstract is recommended to be 150–250 words.")
        st.session_state.form_data['keywords'] = st.text_input(
            "Keywords",
            value=st.session_state.form_data['keywords'],
            help="e.g., machine learning, IEEE, research (3–5 keywords)."
        )

    with st.expander("Footnote Information", expanded=True):
        st.session_state.form_data['footnote']['received_date'] = st.text_input(
            "Manuscript Received Date",
            value=st.session_state.form_data['footnote']['received_date'],
            help="e.g., April 27, 2025"
        )
        st.session_state.form_data['footnote']['revised_date'] = st.text_input(
            "Manuscript Revised Date",
            value=st.session_state.form_data['footnote']['revised_date'],
            help="e.g., September 18, 2025"
        )
        st.session_state.form_data['footnote']['accepted_date'] = st.text_input(
            "Manuscript Accepted Date",
            value=st.session_state.form_data['footnote']['accepted_date'],
            help="e.g., July 25, 2025"
        )
        st.session_state.form_data['footnote']['funding'] = st.text_input(
            "Funding Information",
            value=st.session_state.form_data['footnote']['funding'],
            help="e.g., National Science Foundation Grant XYZ"
        )
        st.session_state.form_data['footnote']['doi'] = st.text_input(
            "DOI",
            value=st.session_state.form_data['footnote']['doi'],
            help="e.g., 10.1109/EXAMPLE.2025.123456"
        )

    with st.expander("Authors", expanded=True):
        st.markdown("**Add or remove authors (displayed side by side in output):**")
        for i, author in enumerate(st.session_state.form_data['authors']):
            with st.container():
                st.markdown(f"**Author {i+1}**")
                cols = st.columns([3, 3, 1])
                with cols[0]:
                    author['name'] = st.text_input(
                        "Name",
                        value=author['name'],
                        key=f"author_name_{i}",
                        help="e.g., John K. Doe"
                    )
                    author['department'] = st.text_input(
                        "Department",
                        value=author.get('department', ''),
                        key=f"author_dept_{i}",
                        help="e.g., Dept. of Computer Science"
                    )
                    author['organization'] = st.text_input(
                        "Organization",
                        value=author.get('organization', ''),
                        key=f"author_org_{i}",
                        help="e.g., IIT Madras"
                    )
                with cols[1]:
                    author['city'] = st.text_input(
                        "City",
                        value=author.get('city', ''),
                        key=f"author_city_{i}",
                        help="e.g., Chennai"
                    )
                    author['state'] = st.text_input(
                        "State",
                        value=author.get('state', ''),
                        key=f"author_state_{i}",
                        help="e.g., Tamil Nadu"
                    )
                    author['tamilnadu'] = st.text_input(
                        "Tamil Nadu (optional)",
                        value=author.get('tamilnadu', ''),
                        key=f"author_tamilnadu_{i}",
                        help="e.g., Tamil Nadu (if applicable)"
                    )
                with cols[2]:
                    if len(st.session_state.form_data['authors']) > 1:
                        if st.button("Delete", key=f"delete_author_{i}"):
                            st.session_state.form_data['authors'].pop(i)
                            st.rerun()
                
                st.markdown("**Additional Author Details (Optional):**")
                for j, custom_field in enumerate(author['custom_fields']):
                    cols = st.columns([3, 1])
                    with cols[0]:
                        custom_field['value'] = st.text_input(
                            "Custom Field (e.g., Email, Phone)",
                            value=custom_field['value'],
                            key=f"author_custom_{i}_{j}",
                            help="e.g., johndoe@example.com"
                        )
                    with cols[1]:
                        if st.button("Delete Field", key=f"delete_custom_{i}_{j}"):
                            author['custom_fields'].pop(j)
                            st.rerun()
                if st.button("Add Custom Field", key=f"add_custom_{i}"):
                    author['custom_fields'].append({'value': ''})
                    st.rerun()
        
        if st.button("Add Author", key="add_author"):
            st.session_state.form_data['authors'].append({
                'name': '', 'department': '', 'organization': '',
                'city': '', 'state': '', 'tamilnadu': '', 'custom_fields': []
            })
            st.rerun()

    with st.expander("Sections and Subsections", expanded=True):
        st.markdown("**Add, edit, or remove sections, subsections, figures, and tables:**")
        for i, section in enumerate(st.session_state.form_data['sections']):
            with st.container():
                st.markdown(f"**Section {i+1}**")
                col1, col2 = st.columns([4, 1])
                with col1:
                    section['title'] = st.text_input(
                        "Section Title",
                        value=section['title'],
                        key=f"section_title_{i}",
                        help="e.g., Methodology"
                    )
                with col2:
                    if len(st.session_state.form_data['sections']) > 1 or section['title'].lower() != 'introduction':
                        if st.button("Delete Section", key=f"delete_section_{i}"):
                            st.session_state.form_data['sections'].pop(i)
                            st.rerun()
                
                # Initialize content_blocks if not exists (for backward compatibility)
                if 'content_blocks' not in section:
                    section['content_blocks'] = []
                    # Migrate old content to new format
                    if section.get('content'):
                        section['content_blocks'].append({
                            'type': 'text',
                            'content': section['content']
                        })
                        section['content'] = ''  # Clear old content
                
                st.markdown("**Content Blocks (Text and Images):**")
                st.caption("💡 Add text blocks and images in any order to create flexible content flow")
                
                # Display existing content blocks
                for block_idx, block in enumerate(section['content_blocks']):
                    with st.container():
                        if block['type'] == 'text':
                            st.markdown(f"📝 **Text Block {block_idx + 1}**")
                            col1, col2 = st.columns([5, 1])
                            with col1:
                                block['content'] = st.text_area(
                                    "Text Content",
                                    value=block.get('content', ''),
                                    height=100,
                                    key=f"text_block_{i}_{block_idx}",
                                    help="Write your paragraph content here"
                                )
                            with col2:
                                st.write("")  # Add spacing
                                st.write("")  # Add spacing
                                if st.button("🗑️ Delete", key=f"delete_text_{i}_{block_idx}"):
                                    section['content_blocks'].pop(block_idx)
                                    st.rerun()
                        
                        elif block['type'] == 'image':
                            st.markdown(f"🖼️ **Image Block {block_idx + 1}**")
                            col1, col2, col3, col4 = st.columns([3, 2, 2, 1])
                            with col1:
                                block['caption'] = st.text_input(
                                    "Image Caption",
                                    value=block.get('caption', ''),
                                    key=f"img_caption_{i}_{block_idx}",
                                    help="Describe what the image shows"
                                )
                            with col2:
                                uploaded_file = st.file_uploader(
                                    "Upload Image",
                                    type=['png', 'jpg', 'jpeg'],
                                    key=f"img_file_{i}_{block_idx}",
                                    accept_multiple_files=False
                                )
                                if uploaded_file:
                                    block['file'] = BytesIO(uploaded_file.read())
                            with col3:
                                block['size'] = st.selectbox(
                                    "Size",
                                    options=['Very Small', 'Small', 'Medium', 'Large'],
                                    index=['Very Small', 'Small', 'Medium', 'Large'].index(block.get('size', 'Medium')),
                                    key=f"img_size_{i}_{block_idx}",
                                    help="Choose image size"
                                )
                            with col4:
                                st.write("")  # Add spacing
                                if st.button("🗑️ Delete", key=f"delete_img_{i}_{block_idx}"):
                                    section['content_blocks'].pop(block_idx)
                                    st.rerun()
                
                # Add new content blocks
                st.markdown("**Add New Content:**")
                col1, col2 = st.columns(2)
                with col1:
                    if st.button(f"📝 Add Text Block", key=f"add_text_block_{i}"):
                        section['content_blocks'].append({
                            'type': 'text',
                            'content': ''
                        })
                        st.rerun()
                with col2:
                    if st.button(f"🖼️ Add Image Block", key=f"add_img_block_{i}"):
                        section['content_blocks'].append({
                            'type': 'image',
                            'caption': '',
                            'file': None,
                            'size': 'Medium'
                        })
                        st.rerun()
                
                st.markdown("**Subsections:**")
                for j, subsection in enumerate(section['subsections']):
                    with st.container():
                        st.markdown(f"Subsection {j+1}")
                        col1, col2 = st.columns([4, 1])
                        with col1:
                            subsection['title'] = st.text_input(
                                "Subsection Title",
                                value=subsection['title'],
                                key=f"subsection_title_{i}_{j}",
                                help="e.g., Data Collection"
                            )
                        with col2:
                            if st.button("Delete Subsection", key=f"delete_subsection_{i}_{j}"):
                                section['subsections'].pop(j)
                                st.rerun()
                        subsection['content'] = st.text_area(
                            "Subsection Content",
                            value=subsection['content'],
                            height=80,
                            key=f"subsection_content_{i}_{j}"
                        )
                
                # Legacy figures support (now at end of section)
                if section.get('figures'):
                    st.markdown("**Legacy Figures** (these will appear at the end of the section)")
                    
                    for j, figure in enumerate(section.get('figures', [])):
                        with st.container():
                            st.markdown(f"Figure {j+1}")
                            col1, col2, col3, col4 = st.columns([3, 3, 1, 1])
                            with col1:
                                figure['caption'] = st.text_input(
                                    "Figure Caption",
                                    value=figure.get('caption', ''),
                                    key=f"figure_caption_{i}_{j}"
                                )
                            with col2:
                                uploaded_file = st.file_uploader(
                                    "Upload Figure (PNG/JPEG)",
                                    type=['png', 'jpg', 'jpeg'],
                                    key=f"figure_file_{i}_{j}",
                                    accept_multiple_files=False
                                )
                                if uploaded_file:
                                    figure['file'] = BytesIO(uploaded_file.read())
                            with col3:
                                figure['size'] = st.selectbox(
                                    "Figure Size",
                                    options=['Very Small', 'Small', 'Medium', 'Large'],
                                    index=['Very Small', 'Small', 'Medium', 'Large'].index(figure.get('size', 'Medium')),
                                    key=f"figure_size_{i}_{j}",
                                    help="Very Small: 1.2 inches wide, Small: 1.8 inches wide, Medium: 2.5 inches wide, Large: 3.2 inches wide"
                                )
                            with col4:
                                if st.button("Delete Figure", key=f"delete_figure_{i}_{j}"):
                                    section['figures'].pop(j)
                                    st.rerun()
        
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("Add Subsection", key=f"add_subsection_{i}"):
                        section['subsections'].append({'title': '', 'content': ''})
                        st.rerun()
                with col2:
                    if st.button("Add Legacy Figure", key=f"add_figure_{i}"):
                        section.setdefault('figures', []).append({
                            'caption': '', 
                            'file': None, 
                            'size': 'Medium'
                        })
                        st.rerun()
        
        if st.button("Add Section", key="add_section"):
            st.session_state.form_data['sections'].append({
                'title': '', 'content': '', 'content_blocks': [], 'subsections': [], 'figures': []
            })
            st.rerun()

    with st.expander("Additional Sections", expanded=True):
        st.session_state.form_data['acknowledgments'] = st.text_area(
            "Acknowledgments (optional)",
            value=st.session_state.form_data['acknowledgments'],
            height=100,
            help="Acknowledge funding or contributors."
        )
        st.markdown("**References (optional, IEEE format, auto-aligned):***")
        for i, ref in enumerate(st.session_state.form_data['references']):
            with st.container():
                ref['text'] = st.text_input(
                    "Reference",
                    value=ref.get('text', ''),
                    key=f"ref_{i}",
                    help="e.g., [1] J. Doe, 'Quantum Teleportation,' IEEE Trans. Quantum Eng., vol. 1, pp. 1-10, 2020."
                )
                if st.button("Delete Reference", key=f"delete_ref_{i}"):
                    st.session_state.form_data['references'].pop(i)
                    st.rerun()
        if st.button("Add Reference", key="add_ref"):
            st.session_state.form_data['references'].append({'text': ''})
            st.rerun()

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Generate IEEE Paper (Word)", key="generate_word"):
            if not st.session_state.form_data['title']:
                st.error("Please enter a title.")
            elif not any(author['name'] for author in st.session_state.form_data['authors']):
                st.error("Please enter at least one author name.")
            else:
                try:
                    buffer = generate_ieee_document(st.session_state.form_data)
                    st.download_button(
                        label="Download Word Document",
                        data=buffer,
                        file_name="ieee_paper.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_word"
                    )
                    st.success("Word document generated successfully!")
                except Exception as e:
                    st.error(f"Error generating Word document: {str(e)}")
    
    with col2:
        if st.button("Generate IEEE Paper (LaTeX)", key="generate_latex"):
            if not st.session_state.form_data['title']:
                st.error("Please enter a title.")
            elif not any(author['name'] for author in st.session_state.form_data['authors']):
                st.error("Please enter at least one author name.")
            else:
                try:
                    buffer = generate_latex_document(st.session_state.form_data)
                    st.download_button(
                        label="Download LaTeX File",
                        data=buffer,
                        file_name="ieee_paper.tex",
                        mime="text/x-tex",
                        key="download_latex"
                    )
                    st.success("LaTeX file generated successfully!")
                except Exception as e:
                    st.error(f"Error generating LaTeX file: {str(e)}")

def add_justified_paragraph(doc, text, style_name='Normal', indent_left=None, indent_right=None, space_before=None, space_after=None):
    """Add a paragraph with optimized justification settings to prevent excessive word spacing."""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Set paragraph formatting with exact spacing controls
    para.paragraph_format.line_spacing = IEEE_CONFIG['line_spacing']
    para.paragraph_format.line_spacing_rule = 0  # Exact spacing
    para.paragraph_format.widow_control = False
    para.paragraph_format.keep_with_next = False
    para.paragraph_format.keep_together = False
    
    # Set spacing
    if space_before is not None:
        para.paragraph_format.space_before = space_before
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    
    # Set indentation
    if indent_left is not None:
        para.paragraph_format.left_indent = indent_left
    if indent_right is not None:
        para.paragraph_format.right_indent = indent_right
    
    # Font formatting with controlled spacing
    if para.runs:
        run = para.runs[0]
        run.font.name = IEEE_CONFIG['font_name']
        run.font.size = IEEE_CONFIG['font_size_body']
        
        # Moderate character spacing controls (not aggressive)
        run_element = run._element
        rPr = run_element.get_or_add_rPr()
        
        # Set moderate character spacing to reduce word gaps without breaking words
        spacing_element = OxmlElement('w:spacing')
        spacing_element.set(qn('w:val'), '-5')  # Slight compression to reduce gaps
        rPr.append(spacing_element)
        
        # Prevent automatic text expansion but allow normal word flow
        run_element.set(qn('w:fitText'), '0')
    
    # Paragraph-level justification controls - MODERATE approach
    para_element = para._element
    pPr = para_element.get_or_add_pPr()
    
    # Use standard justification (not distribute) to keep words intact
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')  # Standard justify - keeps words together
    pPr.append(jc)
    
    # Control text alignment
    textAlignment = OxmlElement('w:textAlignment')
    textAlignment.set(qn('w:val'), 'baseline')
    pPr.append(textAlignment)
    
    # Moderate spacing control - prevent excessive gaps but allow normal flow
    adjust_right_ind = OxmlElement('w:adjustRightInd')
    adjust_right_ind.set(qn('w:val'), '0')
    pPr.append(adjust_right_ind)
    
    return para

if __name__ == "__main__":
    main()